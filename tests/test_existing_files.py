"""Test existing files handling in batch conversion."""

from __future__ import annotations

import tempfile
from pathlib import Path

from nuoyi.api import convert_directory


def test_skip_existing_files():
    """Test skipping existing output files."""
    with tempfile.TemporaryDirectory() as tmpdir:
        input_dir = Path(tmpdir) / "input"
        output_dir = Path(tmpdir) / "output"
        input_dir.mkdir()
        output_dir.mkdir()

        pdf_file = input_dir / "test.pdf"
        pdf_file.write_text("fake pdf content")

        existing_output = output_dir / "test.md"
        existing_output.write_text("existing markdown")

        result = convert_directory(
            input_dir,
            output_dir=output_dir,
            existing_files="skip",
        )

        assert result.success
        assert result.data["skipped"] >= 1
        assert existing_output.read_text() == "existing markdown"


def test_overwrite_existing_files():
    """Test overwriting existing output files."""
    with tempfile.TemporaryDirectory() as tmpdir:
        input_dir = Path(tmpdir) / "input"
        output_dir = Path(tmpdir) / "output"
        input_dir.mkdir()
        output_dir.mkdir()

        from docx import Document

        docx_file = input_dir / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content for overwrite")
        doc.save(str(docx_file))

        existing_output = output_dir / "test.md"
        existing_output.write_text("old content")

        result = convert_directory(
            input_dir,
            output_dir=output_dir,
            existing_files="overwrite",
        )

        assert result.success
        assert result.data["success"] >= 1
        new_content = existing_output.read_text()
        assert "old content" not in new_content


def test_update_only_newer_files():
    """Test updating only when source is newer."""
    with tempfile.TemporaryDirectory() as tmpdir:
        input_dir = Path(tmpdir) / "input"
        output_dir = Path(tmpdir) / "output"
        input_dir.mkdir()
        output_dir.mkdir()

        from docx import Document

        old_file = input_dir / "old.docx"
        doc = Document()
        doc.add_paragraph("old content")
        doc.save(str(old_file))

        old_output = output_dir / "old.md"
        old_output.write_text("old markdown")
        import time

        time.sleep(0.2)
        old_output.touch()

        new_file = input_dir / "new.docx"
        doc2 = Document()
        doc2.add_paragraph("new content")
        doc2.save(str(new_file))

        result = convert_directory(
            input_dir,
            output_dir=output_dir,
            existing_files="update",
            device="cpu",
        )

        assert result.success
        assert old_output.read_text() == "old markdown"
        assert result.data["skipped"] >= 1


if __name__ == "__main__":
    test_skip_existing_files()
    test_overwrite_existing_files()
    test_update_only_newer_files()
    print("All tests passed!")
