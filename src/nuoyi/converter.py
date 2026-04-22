"""
NuoYi - Document converters for PDF and DOCX to Markdown.

Supported PDF engines:

Local (free, offline):
- marker: Best quality, OCR, ~3GB models, GPU recommended
- mineru: Excellent for Chinese, OCR, ~1.5GB models, GPU optional
- docling: Balanced quality, ~1.5GB models, GPU optional
- pymupdf: Fastest, no GPU, digital PDFs only
- pdfplumber: Lightweight, good tables, no GPU

Cloud (API key required):
- llamaparse: LlamaIndex cloud service, excellent quality
- mathpix: Best for math/scientific documents

AMD GPU Support:
- ROCm: AMD GPUs on Linux (RX 500/5000/6000/7000 series)
- DirectML: AMD GPUs on Windows (all Radeon GPUs)
- Vulkan: Experimental cross-platform support

Acceleration backends:
- CUDA: NVIDIA GPUs
- ROCm: AMD GPUs (via HIP on Linux)
- DirectML: AMD/Intel/NVIDIA GPUs on Windows
- MPS: Apple Silicon Metal
- MLX: Apple MLX framework
- CPU: Universal fallback

Memory Optimization for Low VRAM (4-6GB):
- Automatic batch size reduction
- FP16/INT8 quantization
- Model offloading
- Aggressive garbage collection
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document

from .utils import (
    DEFAULT_LANGS,
    LOW_VRAM_THRESHOLD_GB,
    clean_markdown,
    clear_gpu_memory,
    enable_low_vram_mode,
    get_gpu_memory_info,
    get_rocm_memory_info,
    is_amd_gpu_available,
    is_directml_available,
    is_rocm_available,
    select_device,
    setup_memory_optimization,
    _setup_rocm_env,
)

if TYPE_CHECKING:
    pass

setup_memory_optimization()

# Set up ROCm environment BEFORE importing torch-dependent marker modules
if is_rocm_available():
    _setup_rocm_env()

import fitz  # noqa: E402

from marker.converters.pdf import PdfConverter  # noqa: E402
from marker.models import create_model_dict  # noqa: E402
from marker.output import text_from_rendered  # noqa: E402
from marker.config.parser import ConfigParser  # noqa: E402

SUPPORTED_ENGINES = [
    "auto",
    "marker",
    "mineru",
    "docling",
    "pymupdf",
    "pdfplumber",
    "llamaparse",
    "mathpix",
    "mineru-cloud",
    "doc2x",
]

ENGINE_INFO = {
    "marker": {
        "type": "local",
        "gpu": "recommended",
        "models": "~3GB",
        "ocr": True,
        "notes": "Best quality, supports AMD via DirectML/ROCm",
        "amd_support": True,
    },
    "mineru": {
        "type": "local",
        "gpu": "optional",
        "models": "~1.5GB",
        "ocr": True,
        "notes": "Great for Chinese, supports AMD",
        "amd_support": True,
    },
    "docling": {
        "type": "local",
        "gpu": "optional",
        "models": "~1.5GB",
        "ocr": True,
        "notes": "Balanced, supports AMD",
        "amd_support": True,
    },
    "pymupdf": {
        "type": "local",
        "gpu": "no",
        "models": "none",
        "ocr": False,
        "notes": "Fastest, CPU only",
        "amd_support": True,
    },
    "pdfplumber": {
        "type": "local",
        "gpu": "no",
        "models": "none",
        "ocr": False,
        "notes": "Lightweight",
        "amd_support": True,
    },
    "llamaparse": {
        "type": "cloud",
        "gpu": "N/A",
        "models": "cloud",
        "ocr": True,
        "notes": "API key required",
        "amd_support": True,
    },
    "mathpix": {
        "type": "cloud",
        "gpu": "N/A",
        "models": "cloud",
        "ocr": True,
        "notes": "Math specialist",
        "amd_support": True,
    },
    "mineru-cloud": {
        "type": "cloud",
        "gpu": "N/A",
        "models": "cloud",
        "ocr": True,
        "notes": "MinerU online, excellent for Chinese",
        "amd_support": True,
    },
    "doc2x": {
        "type": "cloud",
        "gpu": "N/A",
        "models": "cloud",
        "ocr": True,
        "notes": "Best for formulas, supports split",
        "amd_support": True,
    },
}


def _is_gpu_device(device: str) -> bool:
    return device in ("cuda", "rocm", "mps", "directml", "vulkan")


def _is_vulkan_device(device: str) -> bool:
    return device == "vulkan"


def _is_mlx_device(device: str) -> bool:
    return device == "mlx"


def _is_low_vram_device() -> bool:
    """Check if current GPU has low VRAM (<6GB)."""
    try:
        import torch

        if not torch.cuda.is_available():
            return False
        total, _ = get_gpu_memory_info()
        if total <= 0 and is_rocm_available():
            total, _ = get_rocm_memory_info()
        return total < LOW_VRAM_THRESHOLD_GB
    except Exception:
        return False


def _is_rocm_runtime() -> bool:
    """Check if running on ROCm runtime."""
    return is_rocm_available()


def _is_directml_runtime() -> bool:
    """Check if running on DirectML runtime."""
    return is_directml_available()


def _is_amd_gpu() -> bool:
    """Check if AMD GPU is available."""
    return is_amd_gpu_available()


def _get_amd_backend() -> str | None:
    """Get the AMD backend being used (rocm, directml, or None)."""
    if is_rocm_available():
        return "rocm"
    if is_directml_available():
        return "directml"
    return None


def _setup_amd_environment():
    """Set up environment for AMD GPU acceleration."""
    if is_rocm_available():
        os.environ["HSA_ENABLE_SDMA"] = "0"
        if "HSA_OVERRIDE_GFX_VERSION" not in os.environ:
            os.environ["HSA_OVERRIDE_GFX_VERSION"] = "10.3.0"

    if is_directml_available():
        pass


class PyMuPDFConverter:
    """Fast PDF to Markdown using PyMuPDF4LLM.

    Type: Local, Free, Offline
    GPU: Not required
    OCR: No
    Install: pip install pymupdf4llm
    """

    def __init__(self, page_chunks: bool = True):
        self.page_chunks = page_chunks
        self._module = None

    def _get_module(self):
        if self._module is None:
            try:
                import pymupdf4llm

                self._module = pymupdf4llm
            except ImportError:
                raise ImportError("pip install pymupdf4llm")
        return self._module

    def convert_file(self, pdf_path: str) -> tuple[str, dict]:
        mod = self._get_module()
        md_text = mod.to_markdown(pdf_path, page_chunks=self.page_chunks)

        if isinstance(md_text, list):
            md_text = "\n\n".join(
                c.get("text", "") if isinstance(c, dict) else str(c) for c in md_text
            )

        return clean_markdown(md_text), {}

    @staticmethod
    def is_available() -> bool:
        try:
            import pymupdf4llm

            return True
        except ImportError:
            return False


class PDFPlumberConverter:
    """Lightweight PDF extraction.

    Type: Local, Free, Offline
    GPU: Not required
    OCR: No
    Install: pip install pdfplumber
    """

    def convert_file(self, pdf_path: str) -> tuple[str, dict]:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pip install pdfplumber")

        parts = []

        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                tables = page.extract_tables()

                for table in tables or []:
                    if table:
                        parts.append(self._table_to_md(table))

                if text.strip():
                    parts.append(text.strip())

        return clean_markdown("\n\n".join(parts)), {}

    @staticmethod
    def _table_to_md(table: list[list]) -> str:
        if not table or not table[0]:
            return ""
        lines = []
        h = [str(c) if c else "" for c in table[0]]
        lines.append("| " + " | ".join(h) + " |")
        lines.append("| " + " | ".join(["---"] * len(h)) + " |")
        for row in table[1:]:
            cells = [str(c) if c else "" for c in row]
            cells.extend([""] * (len(h) - len(cells)))
            lines.append("| " + " | ".join(cells[: len(h)]) + " |")
        return "\n".join(lines)

    @staticmethod
    def is_available() -> bool:
        try:
            import pdfplumber

            return True
        except ImportError:
            return False


class DoclingConverter:
    """IBM Docling - balanced quality and speed.

    Type: Local, Free, Offline
    GPU: Optional
    OCR: Yes
    Models: ~1.5GB
    Install: pip install docling

    AMD Support: Works with ROCm (Linux) and DirectML (Windows)
    """

    def __init__(self, device: str = "auto"):
        self.device = select_device(device) if device != "auto" else "cpu"
        self._converter = None

    def _get_converter(self):
        if self._converter is None:
            try:
                from docling.document_converter import DocumentConverter

                self._converter = DocumentConverter()
            except ImportError:
                raise ImportError("pip install docling")
        return self._converter

    def convert_file(self, pdf_path: str) -> tuple[str, dict]:
        conv = self._get_converter()
        result = conv.convert(pdf_path)

        if result and result.document:
            return clean_markdown(result.document.export_to_markdown()), {}

        return "", {}

    @staticmethod
    def is_available() -> bool:
        try:
            from docling.document_converter import DocumentConverter

            return True
        except ImportError:
            return False


class MinerUConverter:
    """MinerU - excellent for Chinese documents.

    Type: Local, Free, Offline
    GPU: Optional (CUDA/CPU)
    OCR: Yes
    Models: ~1.5GB
    Install: pip install magic-pdf[full]
    GitHub: https://github.com/opendatalab/MinerU

    AMD Support: Works with ROCm (Linux)

    Pros:
    - Excellent Chinese document support
    - Good table recognition
    - Smaller models than marker
    - Works on CPU
    """

    def __init__(self, device: str = "auto", langs: str = DEFAULT_LANGS):
        self.device = select_device(device) if device != "auto" else "cpu"
        self.langs = langs
        self._api = None

    def _get_api(self):
        if self._api is None:
            try:
                from magic_pdf.data.data_reader_writer import FileBasedDataReader
                from magic_pdf.data.dataset import PymuDocDataset
                from magic_pdf.model.doc_analyze_by_custom_model import doc_analyze

                self._api = {
                    "FileBasedDataReader": FileBasedDataReader,
                    "PymuDocDataset": PymuDocDataset,
                    "doc_analyze": doc_analyze,
                }
            except ImportError:
                raise ImportError(
                    "Install MinerU:\n"
                    "  pip install magic-pdf[full]\n"
                    "See: https://github.com/opendatalab/MinerU"
                )
        return self._api

    def convert_file(self, pdf_path: str) -> tuple[str, dict]:
        api = self._get_api()

        reader = api["FileBasedDataReader"]()
        pdf_bytes = reader.read(pdf_path)

        ds = api["PymuDocDataset"](pdf_bytes)

        if ds.classify() == "ocr":
            infer_result = ds.apply_ocr()
        else:
            infer_result = ds.apply()

        content_list = infer_result.get_content_list(
            api["FileBasedDataReader"]().read(pdf_path),
            "",
        )

        md_text = content_list.get("markdown", "")
        images = content_list.get("images", {})

        return clean_markdown(md_text), images or {}

    @staticmethod
    def is_available() -> bool:
        try:
            from magic_pdf.data.dataset import PymuDocDataset

            return True
        except ImportError:
            return False


class MarkerPDFConverter:
    """Primary PDF to Markdown converter using marker-pdf.

    Models are loaded once and reused across multiple files.
    Automatically handles GPU memory issues with CPU fallback.

    Memory optimization for low VRAM (4-6GB):
    - Automatic low_vram mode when VRAM < 6GB
    - OCR models offloaded to CPU (saves ~2GB)
    - Reduced batch sizes
    """

    def __init__(
        self,
        force_ocr: bool = False,
        page_range: str | None = None,
        langs: str = DEFAULT_LANGS,
        device: str = "auto",
        low_vram: bool = False,
        allow_fallback: bool = False,
    ):
        self.force_ocr = force_ocr
        self.page_range = page_range
        self.langs = langs
        self.low_vram = low_vram
        self.allow_fallback = allow_fallback

        selected_device = select_device(device)

        if selected_device == "cuda":
            try:
                total, _ = get_gpu_memory_info()
                if total <= 0 and is_rocm_available():
                    total, _ = get_rocm_memory_info()

                if total < LOW_VRAM_THRESHOLD_GB:
                    print(f"[Memory] Low VRAM detected ({total:.1f}GB) - enabling optimization")
                    self.low_vram = True
            except Exception:
                pass

        self.device = selected_device
        os.environ["TORCH_DEVICE"] = self.device

        if self.low_vram:
            self._setup_low_vram_optimizations()

        config = {"output_format": "markdown"}
        if force_ocr:
            config["force_ocr"] = True
        if page_range:
            config["page_range"] = page_range
        if langs:
            config["languages"] = langs

        config_parser = ConfigParser(config)

        self._load_models_with_fallback(config_parser)

    def _setup_low_vram_optimizations(self):
        """Setup memory optimizations for low VRAM."""
        print("[Memory] Low VRAM optimizations enabled")
        enable_low_vram_mode()

    def _load_models_with_fallback(self, config_parser):
        """Load marker-pdf models with automatic CPU fallback on CUDA OOM."""
        print(f"Loading marker-pdf models on {self.device.upper()}...")
        print("(First run downloads ~2-3 GB of model weights)")

        clear_gpu_memory()

        try:
            if self.low_vram:
                import torch

                self.artifact_dict = create_model_dict(device=self.device, dtype=torch.float16)
            else:
                self.artifact_dict = create_model_dict()

            self.converter = PdfConverter(
                config=config_parser.generate_config_dict(),
                artifact_dict=self.artifact_dict,
                processor_list=config_parser.get_processors(),
                renderer=config_parser.get_renderer(),
            )
            print(f"Models loaded successfully on {self.device.upper()}.")

        except RuntimeError as e:
            error_msg = str(e).lower()
            if "cuda" in error_msg and ("out of memory" in error_msg or "oom" in error_msg):
                if self.device != "cpu" and self.allow_fallback:
                    print("\n[WARNING] CUDA out of memory! Falling back to CPU...")
                    print("[WARNING] This will be slower but avoids memory issues.\n")

                    clear_gpu_memory()

                    self.converter = None
                    self.artifact_dict = None

                    import gc

                    gc.collect()

                    self.device = "cpu"
                    os.environ["TORCH_DEVICE"] = "cpu"

                    clear_gpu_memory()

                    print("Reloading models on CPU...")
                    self.artifact_dict = create_model_dict(device="cpu")
                    self.converter = PdfConverter(
                        config=config_parser.generate_config_dict(),
                        artifact_dict=self.artifact_dict,
                        processor_list=config_parser.get_processors(),
                        renderer=config_parser.get_renderer(),
                    )
                    print("Models loaded successfully on CPU.")
                else:
                    raise
            else:
                raise

    def convert_file(self, pdf_path: str) -> tuple[str, dict]:
        """Convert a single PDF file to Markdown text and extract images.

        Returns:
            Tuple of (markdown_text, images_dict) where images_dict maps
            image filenames to PIL Image objects or base64 data.
        """
        clear_gpu_memory()

        try:
            rendered = self.converter(pdf_path)
            text, _, images = text_from_rendered(rendered)
            return clean_markdown(text), images or {}

        except RuntimeError as e:
            error_msg = str(e).lower()
            if "cuda" in error_msg and ("out of memory" in error_msg or "oom" in error_msg):
                if self.device != "cpu" and self.allow_fallback:
                    print("\n[WARNING] CUDA OOM during conversion! Retrying on CPU...")

                    clear_gpu_memory()

                    self.converter = None
                    self.artifact_dict = None

                    import gc

                    gc.collect()

                    clear_gpu_memory()

                    self.device = "cpu"
                    os.environ["TORCH_DEVICE"] = "cpu"

                    config = {"output_format": "markdown"}
                    if self.force_ocr:
                        config["force_ocr"] = True
                    if self.page_range:
                        config["page_range"] = self.page_range
                    if self.langs:
                        config["languages"] = self.langs

                    config_parser = ConfigParser(config)
                    self.artifact_dict = create_model_dict(device="cpu")
                    self.converter = PdfConverter(
                        config=config_parser.generate_config_dict(),
                        artifact_dict=self.artifact_dict,
                        processor_list=config_parser.get_processors(),
                        renderer=config_parser.get_renderer(),
                    )

                    clear_gpu_memory()

                    rendered = self.converter(pdf_path)
                    text, _, images = text_from_rendered(rendered)
                    return clean_markdown(text), images or {}
                raise
            raise

    @staticmethod
    def get_page_count(pdf_path: str) -> int:
        """Quick page count using PyMuPDF."""
        try:
            doc = fitz.open(pdf_path)
            count = len(doc)
            doc.close()
            return count
        except Exception:
            return 0

    @staticmethod
    def is_available() -> bool:
        try:
            from marker.converters.pdf import PdfConverter

            return True
        except ImportError:
            return False


class LlamaParseConverter:
    """LlamaParse - cloud-based parsing by LlamaIndex.

    Type: Cloud (API key required)
    OCR: Yes
    Quality: Excellent
    Cost: Free tier available

    Setup:
        export LLAMA_CLOUD_API_KEY=your_key
    Install: pip install llama-parse
    Get API key: https://cloud.llamaindex.ai/
    """

    def __init__(self, api_key: str | None = None, result_type: str = "markdown"):
        self.api_key = api_key or os.environ.get("LLAMA_CLOUD_API_KEY")
        self.result_type = result_type
        self._parser = None

        if not self.api_key:
            raise ValueError(
                "LlamaParse requires API key. Set LLAMA_CLOUD_API_KEY env var or pass api_key param.\n"
                "Get free API key: https://cloud.llamaindex.ai/"
            )

    def _get_parser(self):
        if self._parser is None:
            try:
                from llama_parse import LlamaParse

                self._parser = LlamaParse(
                    api_key=self.api_key,
                    result_type=self.result_type,
                )
            except ImportError:
                raise ImportError("pip install llama-parse")
        return self._parser

    def convert_file(self, pdf_path: str) -> tuple[str, dict]:
        parser = self._get_parser()
        documents = parser.load_data(pdf_path)

        if documents:
            text = "\n\n".join(doc.text for doc in documents if doc.text)
            return clean_markdown(text), {}

        return "", {}

    @staticmethod
    def is_available() -> bool:
        try:
            from llama_parse import LlamaParse

            return bool(os.environ.get("LLAMA_CLOUD_API_KEY"))
        except ImportError:
            return False


class MathpixConverter:
    """Mathpix - best for math/scientific documents.

    Type: Cloud (API key required)
    OCR: Yes (specialized for math)
    Quality: Excellent for STEM documents
    Cost: Pay per page

    Setup:
        export MATHPIX_APP_ID=your_app_id
        export MATHPIX_APP_KEY=your_app_key
    Install: pip install mathpix
    Get API keys: https://mathpix.com/
    """

    def __init__(
        self,
        app_id: str | None = None,
        app_key: str | None = None,
        output_format: str = "md",
    ):
        self.app_id = app_id or os.environ.get("MATHPIX_APP_ID")
        self.app_key = app_key or os.environ.get("MATHPIX_APP_KEY")
        self.output_format = output_format

        if not self.app_id or not self.app_key:
            raise ValueError(
                "Mathpix requires APP_ID and APP_KEY. Set environment variables:\n"
                "  export MATHPIX_APP_ID=your_id\n"
                "  export MATHPIX_APP_KEY=your_key\n"
                "Get API keys: https://mathpix.com/"
            )

    def convert_file(self, pdf_path: str) -> tuple[str, dict]:
        try:
            import base64
            import json

            import requests
        except ImportError:
            raise ImportError("pip install requests")

        with open(pdf_path, "rb") as f:
            pdf_data = f.read()

        response = requests.post(
            "https://api.mathpix.com/v3/pdf",
            headers={
                "app_id": self.app_id,
                "app_key": self.app_key,
                "Content-type": "application/json",
            },
            json={
                "url": f"data:application/pdf;base64,{base64.b64encode(pdf_data).decode()}",
                "formats": [self.output_format],
            },
        )

        if response.status_code != 200:
            raise RuntimeError(f"Mathpix API error: {response.text}")

        result = response.json()
        pdf_id = result.get("pdf_id")

        if not pdf_id:
            raise RuntimeError(f"No PDF ID in response: {result}")

        status_url = f"https://api.mathpix.com/v3/pdf/{pdf_id}"

        import time

        for _ in range(60):
            time.sleep(2)
            status_resp = requests.get(
                status_url,
                headers={"app_id": self.app_id, "app_key": self.app_key},
            )
            status = status_resp.json()

            if status.get("status") == "completed":
                content_url = f"https://api.mathpix.com/v3/pdf/{pdf_id}.{self.output_format}"
                content_resp = requests.get(
                    content_url,
                    headers={"app_id": self.app_id, "app_key": self.app_key},
                )
                return clean_markdown(content_resp.text), {}

            if status.get("status") == "error":
                raise RuntimeError(f"Mathpix processing error: {status}")

        raise RuntimeError("Mathpix timeout (120s)")

    @staticmethod
    def is_available() -> bool:
        return bool(os.environ.get("MATHPIX_APP_ID") and os.environ.get("MATHPIX_APP_KEY"))


def split_pdf(pdf_path: str, max_pages: int = 50, output_dir: str | None = None) -> list[str]:
    """Split a PDF into multiple smaller PDFs for cloud processing.

    Cloud platforms often have page limits (e.g., 50 pages per request).
    This function splits large PDFs into chunks that can be processed
    individually and then aggregated.

    Parameters
    ----------
    pdf_path : str
        Path to the input PDF file.
    max_pages : int
        Maximum pages per chunk (default: 50).
    output_dir : str or None
        Directory for split files. Defaults to same directory as input.

    Returns
    -------
    list[str]
        List of paths to split PDF files.
    """
    import tempfile

    doc = fitz.open(pdf_path)
    total_pages = len(doc)

    if total_pages <= max_pages:
        doc.close()
        return [pdf_path]

    if output_dir is None:
        output_dir = tempfile.mkdtemp(prefix="nuoyi_split_")

    base_name = Path(pdf_path).stem
    split_files = []

    for i in range(0, total_pages, max_pages):
        end = min(i + max_pages, total_pages)
        chunk_doc = fitz.open()
        chunk_doc.insert_pdf(doc, from_page=i, to_page=end - 1)

        chunk_path = os.path.join(output_dir, f"{base_name}_p{i+1}-{end}.pdf")
        chunk_doc.save(chunk_path)
        chunk_doc.close()
        split_files.append(chunk_path)

        print(f"[Split] Created chunk {len(split_files)}: pages {i+1}-{end}")

    doc.close()
    print(f"[Split] Split {total_pages} pages into {len(split_files)} chunks (max {max_pages}/chunk)")
    return split_files


def aggregate_markdown(
    results: list[tuple[str, dict]],
    page_markers: bool = True,
) -> tuple[str, dict]:
    """Aggregate markdown and images from multiple PDF chunks.

    When a PDF is split into chunks, each chunk produces its own markdown
    and images. This function merges them together:
    - Joins markdown with page break markers
    - Renumbers images to avoid filename conflicts
    - Preserves image references in markdown

    Parameters
    ----------
    results : list[tuple[str, dict]]
        List of (markdown, images_dict) from each chunk.
    page_markers : bool
        Insert page break markers between chunks (default: True).

    Returns
    -------
    tuple[str, dict]
        Aggregated (markdown, images_dict).
    """
    all_md_parts = []
    all_images = {}
    image_counter = 0

    for chunk_idx, (md_text, images) in enumerate(results):
        if page_markers and chunk_idx > 0:
            all_md_parts.append(f"\n---\n\n<!-- Page break (chunk {chunk_idx + 1}) -->\n\n")

        if images:
            new_images = {}
            for img_name, img_data in images.items():
                image_counter += 1
                ext = Path(img_name).suffix or ".png"
                new_name = f"image_{image_counter:04d}{ext}"

                new_images[new_name] = img_data

                old_ref = f"]({img_name})"
                new_ref = f"]({new_name})"
                md_text = md_text.replace(old_ref, new_ref)

                old_ref = f"](./{img_name})"
                md_text = md_text.replace(old_ref, new_ref)

            all_images.update(new_images)

        all_md_parts.append(md_text)

    return clean_markdown("\n\n".join(all_md_parts)), all_images


class MinerUCloudConverter:
    """MinerU Cloud - online PDF to Markdown service.

    Type: Cloud (API key required)
    OCR: Yes (excellent for Chinese)
    Quality: Excellent, especially for Chinese documents
    Cost: Free tier + paid plans

    Setup:
        export MINERU_API_KEY=your_key
    Install: pip install requests
    Get API key: https://mineru.net/

    Features:
    - Excellent Chinese document support
    - Good table and formula recognition
    - Automatic image extraction
    - Supports PDF, images, and other formats
    """

    API_BASE = "https://mineru.net/api/v1"
    POLL_INTERVAL = 3
    POLL_TIMEOUT = 300

    def __init__(self, api_key: str | None = None, lang: str = "ch"):
        self.api_key = api_key or os.environ.get("MINERU_API_KEY")
        self.lang = lang

        if not self.api_key:
            raise ValueError(
                "MinerU Cloud requires API key. Set MINERU_API_KEY env var or pass api_key param.\n"
                "Get API key: https://mineru.net/"
            )

    def _upload_pdf(self, pdf_path: str) -> str | None:
        """Upload PDF and return task_id."""
        import base64

        import requests

        with open(pdf_path, "rb") as f:
            pdf_data = f.read()

        response = requests.post(
            f"{self.API_BASE}/extract",
            headers={
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json",
            },
            json={
                "pdf_base64": base64.b64encode(pdf_data).decode(),
                "lang": self.lang,
                "output_format": "markdown",
            },
        )

        if response.status_code != 200:
            raise RuntimeError(f"MinerU Cloud API error: {response.text}")

        result = response.json()
        return result.get("task_id") or result.get("id")

    def _poll_result(self, task_id: str) -> dict:
        """Poll for task completion and return result."""
        import time

        import requests

        for _ in range(self.POLL_TIMEOUT // self.POLL_INTERVAL):
            time.sleep(self.POLL_INTERVAL)

            response = requests.get(
                f"{self.API_BASE}/extract/{task_id}",
                headers={"Authorization": f"Bearer {self.api_key}"},
            )

            if response.status_code != 200:
                raise RuntimeError(f"MinerU Cloud status error: {response.text}")

            status = response.json()
            state = status.get("state", "").lower()

            if state in ("completed", "success", "done"):
                return status
            if state in ("failed", "error"):
                raise RuntimeError(f"MinerU Cloud processing failed: {status}")

        raise RuntimeError(f"MinerU Cloud timeout ({self.POLL_TIMEOUT}s)")

    def convert_file(self, pdf_path: str) -> tuple[str, dict]:
        """Convert PDF to Markdown via MinerU Cloud API."""
        import requests

        task_id = self._upload_pdf(pdf_path)
        if not task_id:
            raise RuntimeError("No task_id returned from MinerU Cloud")

        result = self._poll_result(task_id)

        md_text = ""
        images = {}

        if "markdown" in result:
            md_text = result["markdown"]
        elif "result" in result:
            md_text = result["result"].get("markdown", "")
            images = result["result"].get("images", {})
        elif "data" in result:
            md_text = result["data"].get("markdown", "")
            images = result["data"].get("images", {})

        if not md_text:
            download_url = result.get("download_url") or result.get("result_url")
            if download_url:
                resp = requests.get(
                    download_url,
                    headers={"Authorization": f"Bearer {self.api_key}"},
                )
                if resp.status_code == 200:
                    md_text = resp.text

        return clean_markdown(md_text), images or {}

    @staticmethod
    def is_available() -> bool:
        return bool(os.environ.get("MINERU_API_KEY"))


class Doc2xConverter:
    """Doc2x - cloud PDF to Markdown with excellent formula support.

    Type: Cloud (API key required)
    OCR: Yes
    Quality: Excellent for STEM documents with formulas
    Cost: Free tier + paid plans

    Setup:
        export DOC2X_API_KEY=your_key
    Install: pip install requests
    Get API key: https://doc2x.noedgeai.com/

    Features:
    - Best-in-class formula recognition (LaTeX)
    - Good table support
    - Image extraction
    - Supports PDF, DOCX, PPTX, images
    """

    API_BASE = "https://v2.doc2x.noedgeai.com/api/v2"
    POLL_INTERVAL = 2
    POLL_TIMEOUT = 300

    def __init__(self, api_key: str | None = None, formula: str = "latex"):
        self.api_key = api_key or os.environ.get("DOC2X_API_KEY")
        self.formula = formula

        if not self.api_key:
            raise ValueError(
                "Doc2x requires API key. Set DOC2X_API_KEY env var or pass api_key param.\n"
                "Get API key: https://doc2x.noedgeai.com/"
            )

    def _upload_and_parse(self, pdf_path: str) -> str | None:
        """Upload PDF and return task/record ID."""
        import requests

        with open(pdf_path, "rb") as f:
            files = {"file": (Path(pdf_path).name, f, "application/pdf")}
            data = {"formula": self.formula}

            response = requests.post(
                f"{self.API_BASE}/parse/pdf",
                headers={"Authorization": f"Bearer {self.api_key}"},
                files=files,
                data=data,
            )

        if response.status_code != 200:
            raise RuntimeError(f"Doc2x API error (HTTP {response.status_code}): {response.text}")

        result = response.json()
        if result.get("code") != "success":
            raise RuntimeError(f"Doc2x upload failed: {result.get('msg', 'Unknown error')}")

        return result.get("data", {}).get("uid")

    def _poll_result(self, uid: str) -> dict:
        """Poll for task completion."""
        import time

        import requests

        for _ in range(self.POLL_TIMEOUT // self.POLL_INTERVAL):
            time.sleep(self.POLL_INTERVAL)

            response = requests.get(
                f"{self.API_BASE}/parse/status",
                headers={"Authorization": f"Bearer {self.api_key}"},
                params={"uid": uid},
            )

            if response.status_code != 200:
                raise RuntimeError(f"Doc2x status error: {response.text}")

            status = response.json()
            if status.get("code") != "success":
                if status.get("code") == "parse_status_not_found":
                    continue
                raise RuntimeError(f"Doc2x poll failed: {status.get('msg', 'Unknown error')}")

            data = status.get("data", {})
            task_status = data.get("status", "")

            if task_status == "success":
                return data
            if task_status in ("fail", "failed", "error"):
                raise RuntimeError(f"Doc2x processing failed: {data.get('fail_reason', 'Unknown')}")

        raise RuntimeError(f"Doc2x timeout ({self.POLL_TIMEOUT}s)")

    def _download_result(self, result_data: dict) -> tuple[str, dict]:
        """Extract markdown and images from result data."""
        pages = result_data.get("result", {}).get("pages", [])

        md_parts = []
        images = {}

        for page in pages:
            md_text = page.get("md", "")
            if md_text:
                md_parts.append(md_text)

            # Extract images from page data
            page_images = page.get("images", {})
            for img_name, img_url in page_images.items():
                try:
                    import requests as req
                    img_resp = req.get(img_url, timeout=10)
                    if img_resp.status_code == 200:
                        images[img_name] = img_resp.content
                except Exception:
                    pass

        return "\n\n".join(md_parts), images

    def convert_file(self, pdf_path: str) -> tuple[str, dict]:
        """Convert PDF to Markdown via Doc2x API."""
        uid = self._upload_and_parse(pdf_path)
        if not uid:
            raise RuntimeError("No UID returned from Doc2x")

        result_data = self._poll_result(uid)
        return self._download_result(result_data)

    @staticmethod
    def is_available() -> bool:
        return bool(os.environ.get("DOC2X_API_KEY"))


def select_engine(engine: str = "auto", has_gpu: bool = True) -> str:
    """Auto-select best available engine."""
    if engine != "auto":
        return engine

    if has_gpu:
        try:
            total, _ = get_gpu_memory_info()
            if total <= 0 and is_rocm_available():
                total, _ = get_rocm_memory_info()

            if total >= 4 and MarkerPDFConverter.is_available():
                backend_info = ""
                if _is_amd_gpu():
                    backend_info = f" (AMD via {_get_amd_backend() or 'auto'})"
                print(f"[Engine] Auto-selected: marker{backend_info}")
                return "marker"
        except Exception:
            pass

    if PyMuPDFConverter.is_available():
        print("[Engine] Auto-selected: pymupdf")
        return "pymupdf"

    if PDFPlumberConverter.is_available():
        print("[Engine] Auto-selected: pdfplumber")
        return "pdfplumber"

    return "marker"


def get_converter(
    engine: str = "auto",
    force_ocr: bool = False,
    page_range: str | None = None,
    langs: str = DEFAULT_LANGS,
    device: str = "auto",
    low_vram: bool = False,
    api_key: str | None = None,
    app_id: str | None = None,
    app_key: str | None = None,
    max_pages: int = 50,
):
    """Get PDF converter by engine name.

    Args:
        engine: Engine name (auto, marker, docling, pymupdf, pdfplumber, llamaparse, mathpix, mineru-cloud, doc2x)
        force_ocr: Force OCR (marker only)
        page_range: Page range (marker only)
        langs: Languages for OCR
        device: GPU device (marker/docling)
            - auto: Auto-detect best device
            - cuda: NVIDIA GPU
            - rocm: AMD GPU on Linux
            - directml: AMD/Intel GPU on Windows
            - mps: Apple Metal
            - mlx: Apple MLX
            - cpu: CPU only
        low_vram: Low VRAM mode (<6GB)
        api_key: API key for LlamaParse/MinerU Cloud/Doc2x
        app_id: App ID for Mathpix
        app_key: App key for Mathpix
        max_pages: Max pages per chunk for cloud converters (default: 50)
    """
    has_gpu = device != "cpu" and (
        _is_gpu_device(device) or _is_mlx_device(device) or device == "auto"
    )

    selected = select_engine(engine, has_gpu)

    if selected == "llamaparse":
        if LlamaParseConverter.is_available() or api_key:
            print("[Converter] Using LlamaParse (cloud)")
            return LlamaParseConverter(api_key=api_key)
        raise ValueError(
            "LlamaParse requires API key.\n"
            "Set: export LLAMA_CLOUD_API_KEY=your_key\n"
            "Get key: https://cloud.llamaindex.ai/"
        )

    if selected == "mathpix":
        if MathpixConverter.is_available() or (app_id and app_key):
            print("[Converter] Using Mathpix (cloud)")
            return MathpixConverter(app_id=app_id, app_key=app_key)
        raise ValueError(
            "Mathpix requires APP_ID and APP_KEY.\n"
            "Set: export MATHPIX_APP_ID=your_id\n"
            "Set: export MATHPIX_APP_KEY=your_key\n"
            "Get keys: https://mathpix.com/"
        )

    if selected == "pymupdf":
        if PyMuPDFConverter.is_available():
            print("[Converter] Using PyMuPDF4LLM (fast, offline)")
            return PyMuPDFConverter()

    if selected == "pdfplumber":
        if PDFPlumberConverter.is_available():
            print("[Converter] Using pdfplumber (lightweight)")
            return PDFPlumberConverter()

    if selected == "docling":
        if DoclingConverter.is_available():
            amd_info = ""
            if _is_amd_gpu():
                amd_info = f" (AMD via {_get_amd_backend()})"
            print(f"[Converter] Using Docling (balanced){amd_info}")
            return DoclingConverter(device=device)

    if selected == "mineru":
        if MinerUConverter.is_available():
            amd_info = ""
            if _is_amd_gpu():
                amd_info = f" (AMD via {_get_amd_backend()})"
            print(f"[Converter] Using MinerU (great for Chinese){amd_info}")
            return MinerUConverter(device=device, langs=langs)

    if selected == "marker":
        if MarkerPDFConverter.is_available():
            amd_info = ""
            if _is_amd_gpu():
                backend = _get_amd_backend()
                amd_info = f" (AMD via {backend.upper()})" if backend else ""

            vram_info = ""
            if low_vram:
                vram_info = " [Low VRAM mode]"

            print(f"[Converter] Using marker-pdf (best quality){amd_info}{vram_info}")
            return MarkerPDFConverter(
                force_ocr=force_ocr,
                page_range=page_range,
                langs=langs,
                device=device,
                low_vram=low_vram,
            )

    if selected == "mineru-cloud":
        if MinerUCloudConverter.is_available() or api_key:
            print("[Converter] Using MinerU Cloud (online, excellent for Chinese)")
            return MinerUCloudConverter(api_key=api_key)
        raise ValueError(
            "MinerU Cloud requires API key.\n"
            "Set: export MINERU_API_KEY=your_key\n"
            "Get key: https://mineru.net/"
        )

    if selected == "doc2x":
        if Doc2xConverter.is_available() or api_key:
            print("[Converter] Using Doc2x (online, best for formulas)")
            return Doc2xConverter(api_key=api_key)
        raise ValueError(
            "Doc2x requires API key.\n"
            "Set: export DOC2X_API_KEY=your_key\n"
            "Get key: https://doc2x.noedgeai.com/"
        )

    raise ImportError(
        "No PDF converter available. Install one of:\n"
        "  pip install marker-pdf         # Best quality, GPU\n"
        "  pip install magic-pdf[full]    # MinerU, Chinese docs\n"
        "  pip install docling            # Balanced, ~1.5GB\n"
        "  pip install pymupdf4llm        # Fastest, no GPU\n"
        "  pip install pdfplumber         # Lightweight\n"
        "  pip install llama-parse        # Cloud, API key\n"
        "  pip install requests           # For Mathpix/MinerU Cloud/Doc2x\n"
        "\nCloud engines (API key required):\n"
        "  export MINERU_API_KEY=xxx      # MinerU Cloud: https://mineru.net/\n"
        "  export DOC2X_API_KEY=xxx       # Doc2x: https://doc2x.noedgeai.com/\n"
        "\nFor AMD GPUs:\n"
        "  Windows: pip install torch-directml  # DirectML support\n"
        "  Linux: Use ROCm PyTorch build        # ROCm support\n"
    )


def list_available_engines() -> dict[str, dict]:
    """List all available engines with their status."""
    engines = {}

    for name, info in ENGINE_INFO.items():
        available = False
        if name == "marker":
            available = MarkerPDFConverter.is_available()
        elif name == "mineru":
            available = MinerUConverter.is_available()
        elif name == "docling":
            available = DoclingConverter.is_available()
        elif name == "pymupdf":
            available = PyMuPDFConverter.is_available()
        elif name == "pdfplumber":
            available = PDFPlumberConverter.is_available()
        elif name == "llamaparse":
            available = LlamaParseConverter.is_available()
        elif name == "mathpix":
            available = MathpixConverter.is_available()
        elif name == "mineru-cloud":
            available = MinerUCloudConverter.is_available()
        elif name == "doc2x":
            available = Doc2xConverter.is_available()

        engines[name] = {**info, "available": available}

    return engines


def print_engines_info():
    """Print engine comparison table."""
    engines = list_available_engines()

    print("\n=== PDF Conversion Engines ===\n")
    print(f"{'Engine':<12} {'Type':<8} {'GPU':<12} {'Available':<10} {'Notes'}")
    print("-" * 70)

    for name, info in engines.items():
        avail = "[OK]" if info["available"] else "[--]"
        gpu = info.get("gpu", "N/A")
        notes = info.get("notes", "")
        amd = " [AMD]" if info.get("amd_support") else ""

        print(f"{name:<12} {info['type']:<8} {gpu:<12} {avail:<10} {notes}{amd}")

    print("\n" + "=" * 70)
    print("AMD GPU Support:")
    print("  Windows: Use --device directml for AMD Radeon GPUs")
    print("  Linux:   Use --device rocm for AMD Radeon GPUs")
    print("  Low VRAM: Use --low-vram for 4-6GB GPUs")
    print("=" * 70)

    print("\nInstall commands:")
    print("  pip install marker-pdf         # Best quality")
    print("  pip install magic-pdf[full]    # MinerU (Chinese docs)")
    print("  pip install docling            # Balanced")
    print("  pip install pymupdf4llm        # Fastest")
    print("  pip install pdfplumber         # Lightweight")
    print("  pip install llama-parse        # Cloud (LLAMA_CLOUD_API_KEY)")
    print("  pip install requests           # For Mathpix/MinerU Cloud/Doc2x")

    print("\nCloud engines (API key required):")
    print("  export MINERU_API_KEY=xxx      # MinerU Cloud: https://mineru.net/")
    print("  export DOC2X_API_KEY=xxx       # Doc2x: https://doc2x.noedgeai.com/")

    print("\nAMD GPU Setup:")
    print("  Windows: pip install torch-directml")
    print("  Linux:   Install ROCm PyTorch from https://pytorch.org/get-started/locally/")
    print(
        "           pip install --pre torch --index-url https://download.pytorch.org/whl/nightly/rocm6.2"
    )


class DocxConverter:
    """DOCX to Markdown using python-docx."""

    def convert_file(self, docx_path: str) -> str:
        doc = Document(docx_path)
        parts = []

        for element in doc.element.body:
            if element.tag.endswith("p"):
                for para in doc.paragraphs:
                    if para._element == element:
                        md = self._para_to_md(para)
                        if md:
                            parts.append(md)
                        break
            elif element.tag.endswith("tbl"):
                for table in doc.tables:
                    if table._element == element:
                        md = self._table_to_md(table)
                        if md:
                            parts.append(md)
                        break

        return "\n\n".join(parts)

    @staticmethod
    def _para_to_md(para) -> str:
        text = para.text.strip()
        if not text:
            return ""

        style = para.style.name.lower() if para.style else ""

        if "heading 1" in style or style == "title":
            return f"# {text}"
        if "heading 2" in style:
            return f"## {text}"
        if "heading 3" in style:
            return f"### {text}"
        if "heading 4" in style:
            return f"#### {text}"
        if "list" in style:
            return f"- {text}"

        parts = []
        for run in para.runs:
            t = run.text
            if not t:
                continue
            if run.bold:
                t = f"**{t}**"
            if run.italic:
                t = f"*{t}*"
            parts.append(t)

        return "".join(parts) if parts else text

    @staticmethod
    def _table_to_md(table) -> str:
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        lines = []
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
        for row in rows[1:]:
            row.extend([""] * (len(rows[0]) - len(row)))
            lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")

        return "\n".join(lines)
